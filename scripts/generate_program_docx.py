import sqlite3
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_document():
    # 1. Connect to Database & Query Programs
    db_path = 'database/database.sqlite'
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Check soft deletes column
    cursor.execute("PRAGMA table_info(programs);")
    columns = [col[1] for col in cursor.fetchall()]
    has_soft_deletes = 'deleted_at' in columns

    if has_soft_deletes:
        cursor.execute("SELECT name FROM programs WHERE deleted_at IS NULL;")
        active_rows = cursor.fetchall()
        cursor.execute("SELECT COUNT(*) FROM programs WHERE deleted_at IS NOT NULL;")
        soft_deleted_count = cursor.fetchone()[0]
    else:
        cursor.execute("SELECT name FROM programs;")
        active_rows = cursor.fetchall()
        soft_deleted_count = 0

    cursor.execute("SELECT COUNT(*) FROM programs;")
    total_count = cursor.fetchone()[0]

    raw_names = [r[0].strip() for r in active_rows if r[0] and r[0].strip()]

    # Turkish Alphabetical Sort Mapping
    tr_map = {
        'a': 'a0', 'A': 'a0', 'b': 'b0', 'B': 'b0', 'c': 'c0', 'C': 'c0',
        'ç': 'c1', 'Ç': 'c1', 'd': 'd0', 'D': 'd0', 'e': 'e0', 'E': 'e0',
        'f': 'f0', 'F': 'f0', 'g': 'g0', 'G': 'g0', 'ğ': 'g1', 'Ğ': 'g1',
        'h': 'h0', 'H': 'h0', 'ı': 'i0', 'I': 'i0', 'i': 'i1', 'İ': 'i1',
        'j': 'j0', 'J': 'j0', 'k': 'k0', 'K': 'k0', 'l': 'l0', 'L': 'l0',
        'm': 'm0', 'M': 'm0', 'n': 'n0', 'N': 'n0', 'o': 'o0', 'O': 'o0',
        'ö': 'o1', 'Ö': 'o1', 'p': 'p0', 'P': 'p0', 'r': 'r0', 'R': 'r0',
        's': 's0', 'S': 's0', 'ş': 's1', 'Ş': 's1', 't': 't0', 'T': 't0',
        'u': 'u0', 'U': 'u0', 'ü': 'u1', 'Ü': 'u1', 'v': 'v0', 'V': 'v0',
        'y': 'y0', 'Y': 'y0', 'z': 'z0', 'Z': 'z0'
    }

    def tr_sort_key(text):
        return ''.join([tr_map.get(ch, ch.lower()) for ch in text])

    # Deduplicate while tracking duplicates
    unique_names_map = {}
    duplicates = []
    for name in raw_names:
        norm = name.lower()
        if norm in unique_names_map:
            duplicates.append(name)
        else:
            unique_names_map[norm] = name

    sorted_unique_names = sorted(unique_names_map.values(), key=tr_sort_key)
    unique_count = len(sorted_unique_names)
    duplicate_count = len(duplicates)

    # 2. Build Word Document
    doc = Document()

    # Set Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set Default Normal Style Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Slate-800

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DOST TV Program Listesi")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B) # Rose-800 / Dark Red Accent

    # Subtitle / Info Line
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tüm Aktif ve Taslak Programlar (Türkçe Alfabetik Sıralı)")
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph() # Spacing

    # Add Program Items (Numbered 1..N)
    for idx, prog_name in enumerate(sorted_unique_names, start=1):
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.line_spacing = 1.15
        
        run_num = p_item.add_run(f"{idx}. ")
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Amber-600

        run_name = p_item.add_run(prog_name)
        run_name.font.color.rgb = RGBColor(0x11, 0x18, 0x27) # Dark slate

    doc.add_paragraph() # Spacing

    # Summary Section
    p_sum_head = doc.add_paragraph()
    p_sum_head.paragraph_format.space_before = Pt(18)
    p_sum_head.paragraph_format.space_after = Pt(6)
    run_sum_head = p_sum_head.add_run("Özet Bilgiler")
    run_sum_head.font.size = Pt(14)
    run_sum_head.font.bold = True
    run_sum_head.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)

    # Summary Bullet Points
    summary_items = [
        ("Toplam program sayısı", str(total_count)),
        ("Tekilleştirilmiş program sayısı", str(unique_count)),
        ("Soft delete kayıt sayısı", str(soft_deleted_count)),
        ("Duplicate kayıt sayısı", str(duplicate_count))
    ]

    for label, val in summary_items:
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(2)
        p_sum.paragraph_format.left_indent = Inches(0.25)
        
        r_bullet = p_sum.add_run("• ")
        r_bullet.font.bold = True
        r_bullet.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

        r_lbl = p_sum.add_run(f"{label}: ")
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        r_val = p_sum.add_run(val)
        r_val.font.bold = True
        r_val.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # 3. Save Output File
    output_filename = "DOST_TV_Program_Listesi.docx"
    doc.save(output_filename)

    # Also save copy to public/exports for web availability
    os.makedirs("public/exports", exist_ok=True)
    doc.save(os.path.join("public/exports", output_filename))

    print(f"Document created successfully: {output_filename}")
    print(f"Total count: {total_count}")
    print(f"Unique count: {unique_count}")
    print(f"Soft deleted count: {soft_deleted_count}")
    print(f"Duplicate count: {duplicate_count}")

if __name__ == '__main__':
    create_document()
